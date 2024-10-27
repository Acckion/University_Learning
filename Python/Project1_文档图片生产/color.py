import json
from docx import Document
from docx.shared import RGBColor

document = Document("img.docx")

with open("color.json") as f:
    colors = json.load(f)
color_idx = 0

for pid, paragraph in enumerate(document.paragraphs):
    original_run = paragraph.runs[0]

    for c in paragraph.text:
        run = paragraph.add_run(c)
        color = colors[color_idx]
        color_idx += 1
        run.font.color.rgb = RGBColor(color[0], color[1], color[2])
        run.font.name = original_run.font.name
        run.font.size = original_run.font.size
    paragraph.runs[0].clear()

document.save("test3.docx")
