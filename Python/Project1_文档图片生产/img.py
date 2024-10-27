import colorsys
import json
from PIL import Image
from docx import Document
from docx.shared import RGBColor


HW_RATIO = 1

img = Image.open("1.jpg")
width, height = img.width, img.height
WIDTH = 90
resized = img.resize((WIDTH, int(height / width * WIDTH * HW_RATIO)))
resized.save("2.jpg")
data = resized.load()

colors = []
print(resized.width, resized.height)
contents = ["肖", "雨", "萍"]

for j in range(resized.height):
    for i in range(resized.width):
        px = data[i, j]
        density = max(px)
        h, s, v = colorsys.rgb_to_hsv(*[d / 255 for d in px])
        s *= 1.2
        s = min(s, 1)
        r, g ,b = colorsys.hsv_to_rgb(h, s, v)
        colors.append([int(255 * r), int(255 * g), int(255 * b)])
        content = contents[(i + j) % len(contents)]
        print(content, end="")
    print("")

with open("color.json", "w") as f:
    json.dump(colors, f)


document = Document("img.docx")

with open("color.json") as f:
    colors = json.load(f)
color_idx = 0

for pid, paragraph in enumerate(document.paragraphs):
    original_run = paragraph.runs[0]

    for c in paragraph.text:
        run = paragraph.add_run(c)
        color = colors[color_idx]
        color_idx += 1
        run.font.color.rgb = RGBColor(color[0], color[1], color[2])
        run.font.name = original_run.font.name
        run.font.size = original_run.font.size
    paragraph.runs[0].clear()

document.save("test3.docx")

