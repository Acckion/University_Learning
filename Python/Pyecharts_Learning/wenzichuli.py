from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image


def create_colored_hanzi_from_image(image_path, output_doc_path, hanzi_list=['肖', '雨', '萍'], font_size=10):
    # 打开图片并转换为RGB模式
    image = Image.open(image_path).convert('RGB')

    # 压缩图片到80x100像素
    image = image.resize((20, 40), Image.Resampling.LANCZOS)

    # 创建Word文档
    doc = Document()

    # 添加一个表格，行数80，列数100
    table = doc.add_table(rows=20, cols=40)

    # 遍历图片的像素点
    for row_index in range(image.height):
        print(row_index)
        for col_index in range(image.width):
            # 获取当前像素的颜色
            pixel_color = image.getpixel((col_index, row_index))
            # 根据像素位置确定使用哪个汉字
            hanzi_index = (row_index * image.width + col_index) % len(hanzi_list)
            hanzi = hanzi_list[hanzi_index]
            # 获取表格单元格
            cell = table.cell(row_index, col_index)
            # 向单元格添加文本
            cell.text = hanzi
            # 设置文本的字体大小和颜色
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor(*pixel_color)

    # 保存文档
    doc.save(output_doc_path)


# 使用示例
image_path = 'D:/Dev/1.jpg'  # 图片路径
output_doc_path = 'D:/Dev/2.docx'  # 输出文档路径
create_colored_hanzi_from_image(image_path, output_doc_path)